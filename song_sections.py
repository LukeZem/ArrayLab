"""Song section detection and Word document export.

Pipeline:
1. Transcribe WAV using Whisper (if available) to get time-stamped segments.
2. Group segments into lyrical blocks using a gap threshold.
3. Identify repeated blocks (chorus) via token Jaccard similarity.
4. Classify remaining blocks as verses or OTHER (bridge / outro etc.). Intro is inferred from initial silence.
5. Export results to a formatted Word document.

The heuristics are intentionally lightweight and deterministic so they can run
offline. They can be swapped for something ML-based later without changing
the public functions.

Public API (stable):
    parse_song_sections(audio_path: str) -> dict
    export_sections_to_docx(sections: dict, output_path: str) -> None

sections dict structure returned by parse_song_sections:
{
  'intro': {'start': float, 'end': float | None},  # intro may have no text
  'verses': [ {'index': int, 'start': float, 'end': float, 'text': str} , ...],
  'choruses': [ {'index': int, 'start': float, 'end': float, 'text': str} , ...],
  'others': [ {'label': 'bridge'|'outro'|'other', 'index': int, 'start': float, 'end': float, 'text': str}, ...],
  'raw_segments': [ Whisper segments list ],
}
"""

from __future__ import annotations

import math
import os
from dataclasses import dataclass
from typing import List, Dict, Any, Iterable

try:
    import whisper  # type: ignore
except ImportError:  # pragma: no cover - environment may not have heavy deps
    whisper = None  # sentinel so we can raise a friendly error later

TOKEN_GAP_SECONDS = 4.0  # gap between segments to start a new block
CHORUS_SIMILARITY_THRESHOLD = 0.6
MIN_TOKENS_FOR_SIMILARITY = 3
MIN_SILENCE_FOR_INTRO = 3.0  # seconds at beginning counts as intro


@dataclass
class Block:
    index: int
    start: float
    end: float
    text: str

    def tokens(self) -> List[str]:
        return [t for t in ''.join(c.lower() if c.isalnum() or c.isspace() else ' ' for c in self.text).split()]


def _load_model(model_size: str = "small"):
    if whisper is None:
        raise RuntimeError(
            "The 'openai-whisper' package (and its deps incl. torch/ffmpeg) is required. "
            "Install via pip and ensure ffmpeg is available."
        )
    return whisper.load_model(model_size)


def transcribe_audio(audio_path: str, model_size: str = "small") -> List[Dict[str, Any]]:
    """Transcribe audio with Whisper returning segments list.

    Each segment: {'id': int, 'seek': int, 'start': float, 'end': float, 'text': str, ...}
    """
    if not os.path.isfile(audio_path):
        raise FileNotFoundError(f"Audio file not found: {audio_path}")
    model = _load_model(model_size)
    result = model.transcribe(audio_path, verbose=False)
    return result.get('segments', [])


def _group_segments_into_blocks(segments: List[Dict[str, Any]]) -> List[Block]:
    blocks: List[Block] = []
    current: List[Dict[str, Any]] = []
    for seg in segments:
        if not current:
            current.append(seg)
            continue
        gap = seg['start'] - current[-1]['end']
        if gap > TOKEN_GAP_SECONDS:
            blocks.append(_segments_to_block(len(blocks), current))
            current = [seg]
        else:
            current.append(seg)
    if current:
        blocks.append(_segments_to_block(len(blocks), current))
    return blocks


def _segments_to_block(index: int, segs: List[Dict[str, Any]]) -> Block:
    text = ' '.join(s['text'].strip() for s in segs).strip()
    return Block(index=index, start=segs[0]['start'], end=segs[-1]['end'], text=text)


def _jaccard(a: Iterable[str], b: Iterable[str]) -> float:
    sa = list(a)
    sb = list(b)
    if not sa or not sb:
        return 0.0
    set_a = set(sa)
    set_b = set(sb)
    inter = len(set_a & set_b)
    denom = min(len(set_a), len(set_b))
    if denom == 0:
        return 0.0
    return inter / denom


def _detect_chorus_blocks(blocks: List[Block]) -> List[int]:
    # Score each block by number of other blocks with high similarity
    similarities = {b.index: [] for b in blocks}
    for i, bi in enumerate(blocks):
        toks_i = bi.tokens()
        if len(toks_i) < MIN_TOKENS_FOR_SIMILARITY:
            continue
        for bj in blocks[i + 1:]:
            toks_j = bj.tokens()
            if len(toks_j) < MIN_TOKENS_FOR_SIMILARITY:
                continue
            sim = _jaccard(toks_i, toks_j)
            if sim >= CHORUS_SIMILARITY_THRESHOLD:
                similarities[bi.index].append(sim)
                similarities[bj.index].append(sim)
    # Choose block with highest count * avg similarity as representative chorus
    chorus_indices: List[int] = []
    scores = []
    for idx, sims in similarities.items():
        if sims:
            scores.append((idx, len(sims), sum(sims) / len(sims)))
    if not scores:
        return []
    scores.sort(key=lambda t: (t[1], t[2]), reverse=True)
    representative = scores[0][0]
    rep_tokens = blocks[representative].tokens()
    for b in blocks:
        sim = _jaccard(rep_tokens, b.tokens())
        if sim >= CHORUS_SIMILARITY_THRESHOLD:
            chorus_indices.append(b.index)
    chorus_indices.sort()
    return chorus_indices


def classify_blocks(blocks: List[Block]) -> Dict[str, Any]:
    if not blocks:
        raise ValueError("No lyrical blocks found; audio may be silent or transcription failed.")
    chorus_indices = _detect_chorus_blocks(blocks)
    verses = []
    others = []
    for b in blocks:
        if b.index in chorus_indices:
            continue
        # Determine if it's a verse: simple heuristic -> earlier blocks that are not chorus
        if not verses:
            verses.append(b)
        else:
            # Compare similarity with first verse; if moderately similar treat as verse
            sim = _jaccard(verses[0].tokens(), b.tokens())
            if sim >= 0.4:
                verses.append(b)
            else:
                label = 'bridge'
                if b.index == len(blocks) - 1:
                    label = 'outro'
                others.append({'label': label, 'index': b.index, 'start': b.start, 'end': b.end, 'text': b.text})
    choruses = [b for b in blocks if b.index in chorus_indices]
    return {
        'blocks': blocks,
        'choruses': [vars(b) for b in choruses],
        'verses': [vars(b) for b in verses],
        'others': others,
    }


def parse_song_sections(audio_path: str, model_size: str = "small") -> Dict[str, Any]:
    """High-level convenience: transcribe and classify sections.
    Returns sections dict described in module docstring."""
    segments = transcribe_audio(audio_path, model_size=model_size)
    if not segments:
        raise ValueError("Transcription produced no segments.")
    intro = None
    first_start = segments[0]['start']
    if first_start >= MIN_SILENCE_FOR_INTRO:
        intro = {'start': 0.0, 'end': first_start, 'text': ''}
    blocks = _group_segments_into_blocks(segments)
    classified = classify_blocks(blocks)
    classified['raw_segments'] = segments
    if intro:
        classified['intro'] = intro
    else:
        classified['intro'] = {'start': 0.0, 'end': blocks[0].start, 'text': ''}
    return classified


def export_sections_to_docx(sections: Dict[str, Any], output_path: str) -> str:
    """Create a formatted Word document summarizing song sections.

    Returns output_path for convenience.
    """
    from docx import Document  # local import keeps module lightweight if not used
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    title = doc.add_heading('Song Structure', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro = sections.get('intro')
    if intro:
        p = doc.add_paragraph()
        p.add_run(f"Intro: 0:00 - {intro['end']:.2f}s\n").bold = True
    # Utility to add section group
    def add_group(name: str, items: List[Dict[str, Any]]):
        if not items:
            return
        doc.add_heading(name, level=1)
        for i, item in enumerate(items, 1):
            runtime = f"[{item['start']:.2f}s - {item['end']:.2f}s]"
            run = doc.add_paragraph().add_run(f"{name[:-1]} {i} {runtime}\n")
            run.bold = True
            lyrics = item['text'].strip()
            if lyrics:
                lyrical_p = doc.add_paragraph(lyrics)
                lyrical_p.paragraph_format.space_after = Pt(12)
    add_group('Choruses', sections.get('choruses', []))
    add_group('Verses', sections.get('verses', []))
    if sections.get('others'):
        doc.add_heading('Other Sections', level=1)
        for other in sections['others']:
            label = other['label'].capitalize()
            runtime = f"[{other['start']:.2f}s - {other['end']:.2f}s]"
            run = doc.add_paragraph().add_run(f"{label} {runtime}\n")
            run.bold = True
            if other['text'].strip():
                doc.add_paragraph(other['text'].strip())
    doc.add_page_break()
    # Raw transcript (optional)
    doc.add_heading('Full Transcript', level=1)
    for seg in sections.get('raw_segments', []):
        doc.add_paragraph(f"{seg['start']:.2f}-{seg['end']:.2f}: {seg['text'].strip()}")
    doc.save(output_path)
    return output_path


def process_song_to_doc(audio_path: str, output_docx: str, model_size: str = "small") -> str:
    """One-shot convenience: parse sections then export the Word document."""
    sections = parse_song_sections(audio_path, model_size=model_size)
    return export_sections_to_docx(sections, output_docx)


if __name__ == "__main__":  # simple CLI
    import argparse
    parser = argparse.ArgumentParser(description="Transcribe a song WAV and export a structured Word doc.")
    parser.add_argument('audio', help='Path to input .wav file')
    parser.add_argument('output', help='Output .docx path')
    parser.add_argument('--model', default='small', help='Whisper model size (tiny, base, small, medium, large)')
    args = parser.parse_args()
    out = process_song_to_doc(args.audio, args.output, model_size=args.model)
    print(f"Created {out}")
